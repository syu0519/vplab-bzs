"""
佐證資料管理系統 - 本機伺服器
執行：python server.py
瀏覽器：http://localhost:8899
"""

import json
import os
import shutil
import io
import base64
import urllib.request
from pathlib import Path
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs

try:
    import fitz  # pymupdf
except ImportError:
    print("""
+----------------------------------------------+
|  ERROR: missing package pymupdf              |
|  Windows:  venv\\Scripts\\activate           |
|            pip install pymupdf               |
|  Mac/Linux: source venv/bin/activate         |
|             pip install pymupdf              |
+----------------------------------------------+
""")
    import sys; sys.exit(1)

# ── 字型自動偵測 ──────────────────────────────
def _find_font(*candidates):
    for p in candidates:
        if Path(p).exists():
            return str(p)
    return None

FONT_CJK = _find_font(
    # macOS
    "/Library/Fonts/MingLiU.ttf",
    "/Library/Fonts/mingliu.ttf",
    # Windows
    r"C:\Windows\Fonts\mingliu.ttc",
    r"C:\Windows\Fonts\MINGLIU.TTC",
    # Linux（fallback）
    "/usr/share/fonts/truetype/arphic/uming.ttc",
)
FONT_LATIN = _find_font(
    # macOS
    "/Library/Fonts/Times New Roman.ttf",
    "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    # Windows
    r"C:\Windows\Fonts\times.ttf",
    r"C:\Windows\Fonts\TIMES.TTF",
)

if FONT_CJK:
    print(f"[字型] CJK  ← {FONT_CJK}")
else:
    print("[字型] CJK 字型未找到，使用內建 china-s（字形較窄）")
if FONT_LATIN:
    print(f"[字型] Latin← {FONT_LATIN}")
else:
    print("[字型] Latin 字型未找到，使用內建 Times-Roman")

# ── 路徑設定 ──
BASE    = Path(__file__).parent
DB_PATH = BASE / "佐證資料庫.json"
UPLOAD  = BASE / "uploads"
OUTPUT  = BASE / "output"
STATIC  = BASE / "static"

UPLOAD.mkdir(exist_ok=True)
OUTPUT.mkdir(exist_ok=True)

# ── AI 端點設定（從 config.json 讀取，預設 localhost）──
_cfg_path = BASE / "config.json"
if _cfg_path.exists():
    _cfg = json.loads(_cfg_path.read_text(encoding="utf-8"))
    LLAMA_URL = _cfg.get("llama_url", "http://localhost:8288")
else:
    LLAMA_URL = "http://localhost:8288"
print(f"[AI端點] {LLAMA_URL}")

# ── 初始化資料庫 ──
INDICATORS = {
    "A1-1":{"section":"壹、教學方面","title":"一、教學內容(10%)","sub":"(一)課程大綱之明確性 15%","note":""},
    "A1-2":{"section":"壹、教學方面","title":"一、教學內容(10%)","sub":"(二)課程教授目標之清楚程度 15%","note":""},
    "A1-3":{"section":"壹、教學方面","title":"一、教學內容(10%)","sub":"(三)教學內容之組織性與系統性 15%","note":""},
    "A1-4":{"section":"壹、教學方面","title":"一、教學內容(10%)","sub":"(四)教學內容之充實性與更新情形 15%","note":""},
    "A1-5":{"section":"壹、教學方面","title":"一、教學內容(10%)","sub":"(五)教學內容與學生需求之配合程度 20%","note":""},
    "A1-6":{"section":"壹、教學方面","title":"一、教學內容(10%)","sub":"(六)教學成果符合授課目標 20%","note":""},
    "A2-1":{"section":"壹、教學方面","title":"二、教學態度(15%)","sub":"(一)教學態度認真、富有熱誠 20%","note":""},
    "A2-2":{"section":"壹、教學方面","title":"二、教學態度(15%)","sub":"(二)授課出勤、缺調補課之情形 20%","note":""},
    "A2-3":{"section":"壹、教學方面","title":"二、教學態度(15%)","sub":"(三)認真批改學生作業並給予學習建議 20%","note":""},
    "A2-4":{"section":"壹、教學方面","title":"二、教學態度(15%)","sub":"(四)能利用課餘時間對學生實施補救教學 20%","note":""},
    "A2-5":{"section":"壹、教學方面","title":"二、教學態度(15%)","sub":"(五)與學生教學互動行為之恰當性 20%","note":""},
    "A3-1":{"section":"壹、教學方面","title":"三、教學方式(15%)","sub":"(一)能清楚表達授課內容 25%","note":""},
    "A3-2":{"section":"壹、教學方面","title":"三、教學方式(15%)","sub":"(二)上課方式有助於引發學生學習動機 25%","note":""},
    "A3-3":{"section":"壹、教學方面","title":"三、教學方式(15%)","sub":"(三)能鼓勵學生發問，尊重其意見表達 25%","note":""},
    "A3-4":{"section":"壹、教學方面","title":"三、教學方式(15%)","sub":"(四)注意教材教法應用，能自製教學媒體配合教學 25%","note":""},
    "A4-1":{"section":"壹、教學方面","title":"四、相關教學行為(20%)","sub":"(一)教師教學與學生相處時行為之恰當性 20%","note":""},
    "A4-2":{"section":"壹、教學方面","title":"四、相關教學行為(20%)","sub":"(二)成績評定方式公平合理 20%","note":""},
    "A4-3":{"section":"壹、教學方面","title":"四、相關教學行為(20%)","sub":"(三)定期積極參與系所教學研討會 20%","note":""},
    "A4-4":{"section":"壹、教學方面","title":"四、相關教學行為(20%)","sub":"(四)授課時教室秩序管理得宜 20%","note":""},
    "A4-5":{"section":"壹、教學方面","title":"四、相關教學行為(20%)","sub":"(五)主動參加各類(校內、外)教學研習活動 20%","note":""},
    "A5-1":{"section":"壹、教學方面","title":"五、教學行政配合(20%)","sub":"(一)教學與校院系所之整體配合情形 40%","note":""},
    "A5-2":{"section":"壹、教學方面","title":"五、教學行政配合(20%)","sub":"(二)教學與教務單位整體配合情形 30%","note":""},
    "A5-3":{"section":"壹、教學方面","title":"五、教學行政配合(20%)","sub":"(三)教學與其他相關單位整體配合情形 30%","note":""},
    "A6-1":{"section":"壹、教學方面","title":"六、教學評量(20%)","sub":"近三年教學評量平均成績 90分以上","note":"（近三年教學評量曾有單一科目低於70分之情形：否。）"},
    "B1-1":{"section":"貳、服務方面","title":"一、兼任行政工作情形(25%)","sub":"各院(系、所、中心)簽請任務指派之行政職務","note":"（本項目不設年限，由教師填寫後，再送人事室查覈。每學期6分。）"},
    "B2-1":{"section":"貳、服務方面","title":"二、全校性服務(25%)","sub":"(一)協助參與學校交辦之工作以配合學校發展 20%","note":""},
    "B2-2":{"section":"貳、服務方面","title":"二、全校性服務(25%)","sub":"(二)參與校級各項委員會情形 20%","note":""},
    "B2-3":{"section":"貳、服務方面","title":"二、全校性服務(25%)","sub":"(三)出缺勤情形 20%","note":""},
    "B2-4":{"section":"貳、服務方面","title":"二、全校性服務(25%)","sub":"(四)參與全校招生工作情形 20%","note":""},
    "B2-5":{"section":"貳、服務方面","title":"二、全校性服務(25%)","sub":"(五)參與學生就業輔導情形 20%","note":""},
    "B3-1":{"section":"貳、服務方面","title":"三、院系所服務(20%)","sub":"(一)協助完成院系所推展工作之情形 30%","note":""},
    "B3-2":{"section":"貳、服務方面","title":"三、院系所服務(20%)","sub":"(二)「招生相關事宜」參與辦理 30%","note":""},
    "B3-3":{"section":"貳、服務方面","title":"三、院系所服務(20%)","sub":"(三)參與系所各項委員會 10%","note":""},
    "B3-4":{"section":"貳、服務方面","title":"三、院系所服務(20%)","sub":"(四)參與系所課程規劃或專業教室規劃及管理 15%","note":""},
    "B3-5":{"section":"貳、服務方面","title":"三、院系所服務(20%)","sub":"(五)輔導學生參加考照、推甄等考試 15%","note":""},
    "B4-1":{"section":"貳、服務方面","title":"四、學輔服務(20%)","sub":"(一)積極參與學生輔導研習會、導師知能研習營或系所導師會議 10%","note":""},
    "B4-2":{"section":"貳、服務方面","title":"四、學輔服務(20%)","sub":"(二)參與學生的班級活動或社團情形 15%","note":""},
    "B4-3":{"section":"貳、服務方面","title":"四、學輔服務(20%)","sub":"(三)擔任導師、社團、學會指導教師、球隊教練或其他校內團隊之指導教師情形 25%","note":""},
    "B4-4":{"section":"貳、服務方面","title":"四、學輔服務(20%)","sub":"(四)對學生輔導工作有具體之行為成效 20%","note":""},
    "B4-5":{"section":"貳、服務方面","title":"四、學輔服務(20%)","sub":"(五)協助學生解決問題，並與學務單位聯繫 30%","note":""},
    "B5-1":{"section":"貳、服務方面","title":"五、推廣服務(5%)","sub":"(一)參與學校推廣班規劃、開發、設立或配合授課 30%","note":""},
    "B5-2":{"section":"貳、服務方面","title":"五、推廣服務(5%)","sub":"(二)配合學校進修部課程教授 30%","note":""},
    "B5-3":{"section":"貳、服務方面","title":"五、推廣服務(5%)","sub":"(三)以學校名義參與相關領域社區或校外服務 40%","note":""},
    "B6-1":{"section":"貳、服務方面","title":"六、服務年資(5%)","sub":"服務年資滿3年基準80分，每增加1年加1分","note":"（本項目不設年限，由教師填寫後，再送人事室查覈。服務年資滿3年之基準分為80分，每增加1年加1分。）"},
}

def load_db():
    if DB_PATH.exists():
        db = json.loads(DB_PATH.read_text(encoding="utf-8"))
        # 自動清理 None items
        for code in db:
            if isinstance(db[code], dict) and "items" in db[code]:
                db[code]["items"] = [it for it in db[code]["items"] if it is not None]
        return db
    # 初始化空資料庫
    db = {}
    for code, info in INDICATORS.items():
        db[code] = {**info, "items": [], "self_score": 0}
    DB_PATH.write_text(json.dumps(db, ensure_ascii=False, indent=2), encoding="utf-8")
    return db

def save_db(db):
    DB_PATH.write_text(json.dumps(db, ensure_ascii=False, indent=2), encoding="utf-8")


# ════════════════════════════════════════
# PDF 蓋章 + 分隔頁（內嵌）
# ════════════════════════════════════════
def img_to_pdf_bytes(img_bytes, ext):
    """把圖片轉成 A4 PDF bytes，自動處理 EXIF 旋轉"""
    # 用 fitz 開圖取得正確方向
    img_doc = fitz.open(stream=img_bytes, filetype=ext if ext != "jpg" else "jpeg")
    pix = img_doc[0].get_pixmap(dpi=150)
    # pix 已套用 EXIF 旋轉
    corrected_bytes = pix.tobytes("jpeg")
    iw, ih = pix.width, pix.height

    # 依圖片方向決定 A4 頁面方向（橫/直）
    if iw > ih:
        pw, ph = 842, 595   # 橫向 A4
    else:
        pw, ph = 595, 842   # 直向 A4

    doc = fitz.open()
    page = doc.new_page(width=pw, height=ph)
    margin = 20
    rect = fitz.Rect(margin, margin, pw - margin, ph - margin)
    page.insert_image(rect, stream=corrected_bytes)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def stamp_pdf_bytes(pdf_bytes, code, start_num=None):
    """在 PDF 每頁視覺右上角蓋字號
    start_num: 若指定，每頁流水號 code(start_num), code(start_num+1)...
               若 None，每頁都蓋同一個 code
    """
    if FONT_CJK:
        kw = {"fontfile": FONT_CJK}
        if FONT_CJK.lower().endswith(".ttc"): kw["fontindex"] = 0
    else:
        kw = {"fontname": "china-s"}

    src = fitz.open("pdf", pdf_bytes)
    out_doc = fitz.open()

    for page_i, src_page in enumerate(src):
        label = f"佐證資料 {code}({start_num+page_i})" if start_num is not None else f"佐證資料 {code}"

        mat = fitz.Matrix(1.5, 1.5)
        pix = src_page.get_pixmap(matrix=mat, alpha=False)
        iw, ih = pix.width, pix.height
        pw, ph = (842, 595) if iw > ih else (595, 842)

        new_page = out_doc.new_page(width=pw, height=ph)
        new_page.insert_image(fitz.Rect(0, 0, pw, ph), pixmap=pix)

        bw = 185 if start_num is not None else 160
        bh = 17
        x0, y0 = pw - bw - 4, 3
        new_page.draw_rect(fitz.Rect(x0, y0, x0+bw, y0+bh),
                           color=NAVY_C, fill=(0.741,0.843,0.933), width=0.8)
        try:
            new_page.insert_text(fitz.Point(x0+6, y0+bh-4),
                                 label, fontsize=9, color=NAVY_C, **kw)
        except Exception:
            new_page.insert_text(fitz.Point(x0+6, y0+bh-4),
                                 label, fontsize=9, color=NAVY_C, fontname="china-s")

    out = io.BytesIO()
    out_doc.save(out)
    return out.getvalue()
    if FONT_CJK:
        kw = {"fontfile": FONT_CJK}
        if FONT_CJK.lower().endswith(".ttc"): kw["fontindex"] = 0
    else:
        kw = {"fontname": "china-s"}

    src = fitz.open("pdf", pdf_bytes)
    out_doc = fitz.open()

    for src_page in src:
        rot = src_page.rotation
        # 把每頁渲染成圖（解決旋轉問題），再放進新 A4 頁
        mat = fitz.Matrix(1.5, 1.5)   # 150% 解析度，清晰度夠
        pix = src_page.get_pixmap(matrix=mat, alpha=False)
        iw, ih = pix.width, pix.height

        # 決定新頁方向
        if iw > ih:
            pw, ph = 842, 595
        else:
            pw, ph = 595, 842

        new_page = out_doc.new_page(width=pw, height=ph)
        new_page.insert_image(fitz.Rect(0, 0, pw, ph), pixmap=pix)

        # 蓋章（座標現在完全正常）
        bw, bh = 160, 17
        x0, y0 = pw - bw - 4, 3
        new_page.draw_rect(fitz.Rect(x0, y0, x0+bw, y0+bh),
                           color=NAVY_C, fill=(0.741,0.843,0.933), width=0.8)
        try:
            new_page.insert_text(fitz.Point(x0+6, y0+bh-4),
                                 f"佐證資料 {code}", fontsize=9, color=NAVY_C, **kw)
        except Exception:
            new_page.insert_text(fitz.Point(x0+6, y0+bh-4),
                                 f"佐證資料 {code}", fontsize=9, color=NAVY_C, fontname="china-s")

    out = io.BytesIO()
    out_doc.save(out)
    return out.getvalue()

NAVY_C = (0.122, 0.306, 0.475)  # 模組層常數，stamp 用


def make_divider_bytes(ind_code, db, expanded=False):
    """產生分隔頁 PDF bytes（支援多行說明、多頁）
    expanded=True：群組展開，每個 item 各自一列（用於精華摘要）
    expanded=False：群組合併一列（預設，用於卷夾分隔頁）
    """
    ind = db.get(ind_code, INDICATORS.get(ind_code, {}))
    items = [it for it in ind.get("items", []) if it is not None]
    W, H = 595, 842
    NAVY   = (0.122, 0.306, 0.475)
    ORANGE = (0.773, 0.353, 0.067)
    WHITE  = (1.0, 1.0, 1.0)
    GRAY_ROW = (0.94, 0.965, 0.988)

    CX = [58,  82, 170, 214, 252]
    CW = [24,  88,  44,  38, 265]
    FSIZ  = 8.5
    RH_MIN = 20
    LINE_H = 11

    def _ins(page, pt, text, size, color):
        """通用插入，自動選字型"""
        has_cjk = any('\u4e00'<=c<='\u9fff' for c in str(text))
        if has_cjk:
            kw = {"fontfile": FONT_CJK} if FONT_CJK else {"fontname": "china-s"}
        else:
            kw = {"fontfile": FONT_LATIN} if FONT_LATIN else {"fontname": "Times-Roman"}
        page.insert_text(pt, str(text), fontsize=size, color=color, **kw)

    def new_page(doc):
        page = doc.new_page(width=W, height=H)
        # 頂部水平線移除，只保留標題文字
        _ins(page, fitz.Point(58, 68),  ind.get("section",""), 15, NAVY)
        _ins(page, fitz.Point(72, 96),  ind.get("title",""),   13, NAVY)
        _ins(page, fitz.Point(86, 116), ind.get("sub",""),     10, (0.35,0.35,0.35))
        note = ind.get("note","")
        if note:
            _ins(page, fitz.Point(86, 134), note, 8.5, (0.55,0.55,0.55))
        return page

    def draw_table_header(page, ty):
        TH = 20
        page.draw_rect(fitz.Rect(CX[0], ty, CX[0]+sum(CW), ty+TH),
                       color=None, fill=NAVY)
        headers = ["編號", "佐證資料編號", "年度", "頁數", "說　　明"]
        offsets = [3, 3, 3, 4, 3]
        for i, h in enumerate(headers):
            _ins(page, fitz.Point(CX[i]+offsets[i], ty+TH-6), h, FSIZ, WHITE)
        return ty + TH

    def draw_footer(page):
        # 底部水平線移除，只保留頁腳文字
        _ins(page, fitz.Point(58, H-28),
            "嶺東科技大學 設計時尚學院 數位媒體設計系　教師資格審查佐證資料",
            8, (0.55,0.55,0.55))
        _ins(page, fitz.Point(W-110, H-28), f"佐證代碼：{ind_code}", 8, NAVY)

    def wrap_text(text, max_chars):
        """把長文字按字數換行"""
        if not text: return [""]
        lines = []
        while len(text) > max_chars:
            lines.append(text[:max_chars])
            text = text[max_chars:]
        if text: lines.append(text)
        return lines

    def _ins(page, pt, text, size, color):
        """通用插入，自動選字型"""
        text = str(text)
        has_cjk = any('\u4e00'<=c<='\u9fff' for c in text)
        try:
            if has_cjk and FONT_CJK:
                kw = {"fontfile": FONT_CJK}
                if FONT_CJK.lower().endswith(".ttc"): kw["fontindex"] = 0
                page.insert_text(pt, text, fontsize=size, color=color, **kw)
            elif not has_cjk and FONT_LATIN:
                page.insert_text(pt, text, fontsize=size, color=color, fontfile=FONT_LATIN)
            elif has_cjk:
                page.insert_text(pt, text, fontsize=size, color=color, fontname="china-s")
            else:
                page.insert_text(pt, text, fontsize=size, color=color, fontname="Times-Roman")
        except Exception:
            # fallback
            fn = "china-s" if has_cjk else "Times-Roman"
            page.insert_text(pt, text, fontsize=size, color=color, fontname=fn)

    def insert_cell(page, pt, text, color=(0.2,0.2,0.2), size=FSIZ, force_cjk=False):
        text = str(text)
        has_cjk = force_cjk or any('\u4e00'<=c<='\u9fff' for c in text)
        try:
            if has_cjk and FONT_CJK:
                kw = {"fontfile": FONT_CJK}
                if FONT_CJK.lower().endswith(".ttc"): kw["fontindex"] = 0
                page.insert_text(pt, text, fontsize=size, color=color, **kw)
            elif not has_cjk and FONT_LATIN:
                page.insert_text(pt, text, fontsize=size, color=color, fontfile=FONT_LATIN)
            elif has_cjk:
                page.insert_text(pt, text, fontsize=size, color=color, fontname="china-s")
            else:
                page.insert_text(pt, text, fontsize=size, color=color, fontname="Times-Roman")
        except Exception:
            fn = "china-s" if has_cjk else "Times-Roman"
            page.insert_text(pt, text, fontsize=size, color=color, fontname=fn)

    # ── 把 items 轉成「顯示列」（群組合併） ──
    # item 若有 group 欄位，同 group 合併成一列
    # 沒有 group 的 item 各自一列
    from collections import OrderedDict

    display_rows = []   # [{code_str, year_str, pages, desc, items:[...]}]
    group_map = OrderedDict()

    for item in items:
        g = (item.get("group") or "").strip()
        if g:
            if g not in group_map:
                group_map[g] = []
            group_map[g].append(item)
        else:
            # 無群組 → 直接加一列
            display_rows.append({
                "code_str": item.get("code",""),
                "year_str": str(item.get("year","")),
                "pages": item.get("pages",1),
                "desc": item.get("desc","").strip() or item.get("ocr","")[:50].strip()
                        or (item.get("file","") or "").replace("\\","/").split("/")[-1],
                "is_group": False,
            })

    # 群組的先依 group key 順序全部加進去（放在最後，或可依需求調整）
    # 但要保持原本 items 順序中 group 出現的位置
    # 重新掃一次，保持順序
    display_rows = []
    seen_groups = set()
    for item in items:
        g = (item.get("group") or "").strip()
        if g:
            if g in seen_groups:
                if expanded:
                    # 展開模式：群組已經輸出標題，直接輸出 item 本身
                    raw_desc = item.get("desc","").strip() or (item.get("file","") or "").split("/")[-1]
                    code = item.get("code","")
                    dtype = item.get("doc_type","") or "佐證文件"
                    pages = item.get("pages",1)
                    skip = int(item.get("skip_pages",0) or 0)
                    real_pages = max(1, pages - skip)
                    code_display = f"{code}(1)~({real_pages})" if real_pages > 1 else code
                    display_rows.append({
                        "code_str": code_display,
                        "year_str": str(item.get("year","")),
                        "pages": pages,
                        "desc": raw_desc + f"\n[佐證資料{code_display}：{dtype}]",
                        "is_group": False,
                        "is_sub": True,  # 縮排標示
                    })
                continue

            seen_groups.add(g)
            grp_items = group_map[g]
            total_pages = sum(it.get("pages",1) for it in grp_items)
            codes = [it.get("code","") for it in grp_items if it.get("code")]
            import re as _re3
            all_year_nums = []
            for it in grp_items:
                y = str(it.get("year","")).strip()
                nums = [int(n) for n in _re3.findall(r'\d{3}', y)]
                all_year_nums.extend(nums)
            if all_year_nums:
                mn, mx = min(all_year_nums), max(all_year_nums)
                year_display = f"{mn}-{mx}" if mn != mx else str(mn)
            else:
                year_display = ""
            group_label = g

            if expanded:
                # 展開模式：先輸出群組標題列，再逐一輸出每個 item
                display_rows.append({
                    "code_str": f"{codes[0]}~{codes[-1]}" if len(codes)>1 else (codes[0] if codes else ""),
                    "year_str": year_display,
                    "pages": total_pages,
                    "desc": f"▌ {group_label}（共 {len(grp_items)} 份）",
                    "is_group": True,
                    "count": len(grp_items),
                })
                for it in grp_items:
                    raw_desc = it.get("desc","").strip() or (it.get("file","") or "").split("/")[-1]
                    code = it.get("code","")
                    dtype = it.get("doc_type","") or "佐證文件"
                    pages_i = it.get("pages",1)
                    skip_i = int(it.get("skip_pages",0) or 0)
                    real_i = max(1, pages_i - skip_i)
                    code_d = f"{code}(1)~({real_i})" if real_i > 1 else code
                    display_rows.append({
                        "code_str": code_d,
                        "year_str": str(it.get("year","")),
                        "pages": pages_i,
                        "desc": f"  • {raw_desc}\n[佐證資料{code_d}：{dtype}]",
                        "is_group": False,
                        "is_sub": True,
                    })
            else:
                # 合併模式（預設）
                if len(grp_items) == 1:
                    it0 = grp_items[0]
                    skip  = int(it0.get("skip_pages",0) or 0)
                    real  = max(1, it0.get("pages",1) - skip)
                    code0 = codes[0] if codes else ""
                    code_display = f"{code0}(1)~({real})" if real > 1 else code0
                    dtype = it0.get("doc_type","") or "佐證文件"
                    desc = f"{group_label}\n[佐證資料{code0}：{dtype}]"
                else:
                    code_display = f"{codes[0]}~{codes[-1]}" if len(codes)>1 else (codes[0] if codes else "")
                    dtypes = [it.get("doc_type","") or "佐證文件" for it in grp_items]
                    dtype = max(set(dtypes), key=dtypes.count)
                    desc = f"{group_label}\n[佐證資料{codes[0]}：{dtype} 共{len(grp_items)}張]"

                display_rows.append({
                    "code_str": code_display,
                    "year_str": year_display,
                    "pages": total_pages,
                    "desc": desc,
                    "is_group": True,
                    "count": len(grp_items),
                })
        else:
            raw_desc = item.get("desc","").strip()
            if not raw_desc:
                raw_desc = item.get("ocr","")[:50].strip()
            if not raw_desc:
                raw_desc = (item.get("file","") or "").replace("\\","/").split("/")[-1]
            pages = item.get("pages",1)
            skip  = int(item.get("skip_pages",0) or 0)
            real_pages = max(1, pages - skip)
            code = item.get("code","")
            dtype = item.get("doc_type","") or "佐證文件"
            if real_pages > 1:
                code_display = f"{code}(1)~({real_pages})"
            else:
                code_display = code
            # 說明欄：desc + [佐證資料code：類型]
            full_desc = raw_desc + f"\n[佐證資料{code_display}：{dtype}]" if raw_desc else f"[佐證資料{code_display}：{dtype}]"
            display_rows.append({
                "code_str": code_display,
                "year_str": str(item.get("year","")),
                "pages": pages,
                "desc": full_desc,
                "is_group": False,
            })

    doc = fitz.open()
    page = new_page(doc)
    TABLE_TOP = 148
    y = draw_table_header(page, TABLE_TOP)
    FOOTER_Y = H - 60   # 低於此 y 就換頁

    for ri, row in enumerate(display_rows):
        # 說明（多行）：先按 \n 分段，再各自 wrap
        raw_desc_text = row["desc"]
        segments = raw_desc_text.split("\n")
        desc_lines = []
        line_colors = []
        for seg in segments:
            wrapped = wrap_text(seg, 30)
            is_code_line = seg.strip().startswith("[佐證資料")
            for wl in wrapped:
                desc_lines.append(wl)
                line_colors.append((
                    ORANGE if is_code_line else (0.1,0.1,0.1),
                    7.5 if is_code_line else FSIZ,
                    True  # force_cjk
                ))
        row_h = max(RH_MIN, len(desc_lines) * LINE_H + 6)

        # 換頁判斷
        if y + row_h > FOOTER_Y:
            draw_footer(page)
            page = new_page(doc)
            y = draw_table_header(page, TABLE_TOP)

        bg = GRAY_ROW if ri % 2 == 0 else WHITE
        if row.get("is_group"):
            bg = (0.878, 0.937, 0.878)
        page.draw_rect(fitz.Rect(CX[0], y, CX[0]+sum(CW), y+row_h),
                       color=(0.8,0.8,0.8), fill=bg, width=0.3)

        # 編號
        cy = y + row_h//2 + 3
        insert_cell(page, fitz.Point(CX[0]+3, cy), f"{ri+1}.")

        # 佐證字號（橘色）
        insert_cell(page, fitz.Point(CX[1]+3, cy),
            row["code_str"], color=ORANGE,
            size=7.5 if row.get("is_group") else FSIZ)

        # 年度
        insert_cell(page, fitz.Point(CX[2]+3, cy), row["year_str"])

        # 頁數
        insert_cell(page, fitz.Point(CX[3]+6, cy), str(row["pages"]))

        # 說明（多行，從頂部開始）
        text_y = y + LINE_H
        for line, (color, size, fcjk) in zip(desc_lines, line_colors):
            insert_cell(page, fitz.Point(CX[4]+3, text_y), line, color=color, size=size, force_cjk=fcjk)
            text_y += LINE_H

        y += row_h

    draw_footer(page)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def build_volume(ind_code, db, include_divider=True):
    """產生完整卷夾 PDF，存到 output/"""
    ind = db.get(ind_code, {})
    items = [it for it in ind.get("items", []) if it is not None]

    print(f"\n[build] {ind_code}，共 {len(items)} 份文件，include_divider={include_divider}")
    print(f"[build] UPLOAD 路徑：{UPLOAD}")

    merged = fitz.open()

    # 1. 分隔頁（可選）
    if include_divider:
        div_bytes = make_divider_bytes(ind_code, db)
        merged.insert_pdf(fitz.open("pdf", div_bytes))

    # 2. 各佐證文件（蓋章）
    IMG_EXTS = {".jpg",".jpeg",".png",".bmp",".gif",".webp",".tiff"}
    for item in items:
        file_rel = item.get("file","")
        print(f"  [item] code={item.get('code')} file={file_rel!r}")
        if not file_rel:
            print(f"  [skip] file 欄位空")
            continue
        pdf_path = UPLOAD / file_rel
        print(f"  [path] {pdf_path}  exists={pdf_path.exists()}")
        if not pdf_path.exists():
            # 嘗試用純檔名在整個 uploads 搜尋
            fname = Path(file_rel).name
            found = list(UPLOAD.rglob(fname))
            if found:
                pdf_path = found[0]
                print(f"  [alt]  找到替代路徑：{pdf_path}")
            else:
                print(f"  [skip] 完全找不到 {fname!r}")
                continue
        ext = pdf_path.suffix.lower()
        try:
            if ext in IMG_EXTS:
                img_bytes = pdf_path.read_bytes()
                pdf_bytes = img_to_pdf_bytes(img_bytes, ext.lstrip("."))
                skip = 0
            else:
                pdf_bytes = pdf_path.read_bytes()
                skip = int(item.get("skip_pages", 0) or 0)

            src_doc = fitz.open("pdf", pdf_bytes)
            total_pages = src_doc.page_count
            src_doc.close()

            if total_pages > 1 or skip > 0:
                # 多頁 PDF：前 skip 頁直接插入（不蓋章），其餘流水號蓋章
                if skip > 0:
                    # 前 skip 頁不蓋章，直接插入
                    front_doc = fitz.open("pdf", pdf_bytes)
                    front_bytes = io.BytesIO()
                    front_out = fitz.open()
                    front_out.insert_pdf(front_doc, from_page=0, to_page=skip-1)
                    front_out.save(front_bytes)
                    merged.insert_pdf(fitz.open("pdf", front_bytes.getvalue()))
                    front_doc.close()

                # 剩餘頁蓋流水號章
                remain_doc = fitz.open("pdf", pdf_bytes)
                remain_bytes = io.BytesIO()
                remain_out = fitz.open()
                remain_out.insert_pdf(remain_doc, from_page=skip, to_page=total_pages-1)
                remain_out.save(remain_bytes)
                remain_out.close()

                stamped = stamp_pdf_bytes(remain_bytes.getvalue(), item["code"], start_num=1)
                merged.insert_pdf(fitz.open("pdf", stamped))
                print(f"  [ok] {pdf_path.name} -> skip={skip} + {item['code']}(1)~({total_pages-skip})")
            else:
                stamped = stamp_pdf_bytes(pdf_bytes, item["code"])
                merged.insert_pdf(fitz.open("pdf", stamped))
                print(f"  [ok] {pdf_path.name} -> {item['code']}")
        except Exception as e:
            print(f"  [err] {pdf_path.name}: {e}")

    suffix = "卷夾" if include_divider else "蓋章"
    out_path = OUTPUT / f"{ind_code}_{suffix}.pdf"
    merged.save(str(out_path))
    merged.close()
    return out_path


# ════════════════════════════════════════
# GPTs Knowledge 產生
# ════════════════════════════════════════

# 各指標建議結尾句
_GPTS_ENDING = {
    "A1-1": "課程大綱教學計劃明確清楚。",
    "A1-2": "課程教授目標明確清楚，具備產業對應性。",
    "A1-3": "教學內容之組織性與系統性完整且良好。",
    "A1-4": "教學內容之充實性與更新情形完整且良好，並持續回應產業技術演進。",
    "A1-5": "教學內容與學生需求之配合良好，並持續更新上傳。",
    "A1-6": "教學成果良好符合授課目標。",
    "A2-1": "教學態度認真、富有熱誠，持續獲推動實務教學獎勵之肯定。",
    "A2-2": "近三年無未依規定辦理調補課、無未按規定繳交成績、無遲到早退或無故缺課。",
    "A2-3": "認真批改學生作業並給予學習建議。",
    "A2-4": "能利用課餘時間對學生實施各種輔導與補救教學。",
    "A2-5": "與學生教學互動行為良好，師生互動成果具體呈現於歷屆專題獲獎與展演紀錄。",
    "A3-1": "教學科目能提供完整之教學計畫，能清楚表達授課內容。",
    "A3-2": "上課方式多元，有助於引發學生學習動機。",
    "A3-3": "能鼓勵學生發問，尊重其意見表達。",
    "A3-4": "注意教材教法應用，能自製教學媒體配合教學。",
    "A4-1": "教師教學與學生相處時行為恰當。",
    "A4-2": "成績評定方式公平合理。",
    "A4-3": "定期積極參與系所及跨校教學研討社群。",
    "A4-4": "授課時教室秩序管理得宜。",
    "A4-5": "主動參加各類校內外教學研習活動。",
    "A5-1": "教學與校院系所整體配合情形良好，持續支援教學、招生與場域服務。",
    "A5-2": "能配合辦理學校重要任務或緊急事務。",
    "A5-3": "教學與校院其他相關單位整體配合情形良好。",
    "B2-1": "持續協助參與學校交辦之工作以配合學校發展。",
    "B2-4": "積極參與全校招生工作，成效良好。",
    "B2-5": "積極參與學生就業輔導，協助學生與業界接軌。",
    "B3-1": "長期協助完成院系所推展工作，貢獻具體且持續。",
    "B3-2": "積極參與院系招生相關事宜，成效顯著。",
    "B3-4": "負責專業教室規劃與管理，維護場域正常運作。",
    "B3-5": "積極輔導學生考照及升學推甄，成效良好。",
    "B4-1": "積極參與學生輔導研習，持續增進導師專業知能。",
    "B4-2": "積極參與班級活動，師生互動良好。",
    "B4-3": "擔任導師及社團指導老師，對學生輔導工作貢獻具體。",
    "B4-4": "對學生輔導工作有具體之行為成效。",
    "B4-5": "積極協助學生解決問題，並與學務單位保持良好聯繫。",
    "B5-3": "積極以學校名義參與相關領域社區及校外服務。",
}

# 需要附評量數據的指標
_NEEDS_SCORE = {"A1-6", "A2-1"}
_SCORE_TEXT = """依近三年（1111～1142）教學統計表：
1111學期平均91.5（採計259人）、1112學期89.5（232人）、
1121學期93.0（240人）、1122學期93.2（224人）、
1131學期92.7（225人）、1132學期91.7（96人）、
1141學期90.9（276人）、1142學期92.2（122人）。"""


def _parse_ai_summary(item):
    """從item中提取AI摘要文字"""
    ai = item.get("ai")
    if not ai or not isinstance(ai, dict):
        return ""
    summary = ai.get("summary", "")
    if not summary:
        return ""
    # 清理 ```json 殘留
    if "```" in summary:
        try:
            clean = summary.replace("```json", "").replace("```", "").strip()
            parsed = json.loads(clean)
            return parsed.get("summary", "")
        except Exception:
            return ""
    return summary[:150]


def build_gpts_knowledge(db):
    """將佐證資料庫json轉成GPTs Knowledge Markdown文字"""
    from datetime import datetime

    lines = []
    lines.append("# 佐證資料庫 Knowledge（自動產生）")
    lines.append(f"# 產生時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 寫作規則提示")
    lines.append("")
    lines.append("- 每格格式：標題行 → 近N年事實條列 → 結尾一句（用「建議結尾句」）")
    lines.append("- 結尾句對應各格指標核心語言，不得通用複製")
    lines.append("- 有「評量數據」欄位的格才引數字，其他格不加")
    lines.append("- 只用本檔案中出現的佐證編號，不憑空補充")
    lines.append("")
    lines.append("---")
    lines.append("")

    for code, ind in db.items():
        if code == "?":
            continue

        section = ind.get("section", "")
        title = ind.get("title", "")
        sub = ind.get("sub", "")
        items = ind.get("items", [])
        self_score = ind.get("self_score", 0)
        note = ind.get("note", "")

        lines.append(f"## 指標 {code}")
        lines.append(f"**章節**：{section}　**大項**：{title}")
        lines.append(f"**細項**：{sub}")
        if self_score:
            lines.append(f"**教師自評**：{self_score} 分")
        if note:
            lines.append(f"**備註**：{note}")
        lines.append("")

        # 結尾句
        ending = _GPTS_ENDING.get(code, "")
        if ending:
            lines.append(f"**建議結尾句**：{ending}")

        # 評量數據
        if code in _NEEDS_SCORE:
            lines.append("")
            lines.append(f"**評量數據**：")
            lines.append(_SCORE_TEXT)

        lines.append("")
        lines.append("**佐證清單**：")
        lines.append("")

        # 依 group 分組
        groups: dict = {}
        for item in items:
            g = item.get("group") or "其他"
            groups.setdefault(g, []).append(item)

        for g_name, g_items in groups.items():
            if g_name != "其他":
                lines.append(f"*【{g_name}】*")
            for item in g_items:
                icode = item.get("code", "")
                year  = item.get("year", "")
                desc  = item.get("desc", "")
                pages = item.get("pages", "")
                status = item.get("status", "")
                doc_type = item.get("doc_type", "")

                mark = "✓" if status == "done" else "○"
                row = f"- {mark} `{icode}` [{year}] {desc}"
                if pages:
                    row += f"（{pages}頁）"
                if doc_type:
                    row += f" [{doc_type}]"
                lines.append(row)

                # AI 摘要
                summary = _parse_ai_summary(item)
                if summary and len(summary) > 10:
                    lines.append(f"  > {summary}")

        lines.append("")
        lines.append("---")
        lines.append("")

    # 統計
    total = sum(len(v.get("items", [])) for k, v in db.items() if k != "?")
    done  = sum(1 for k, v in db.items() if k != "?"
                for item in v.get("items", []) if item.get("status") == "done")
    lines.append(f"## 統計")
    lines.append(f"- 指標總數：{len([k for k in db if k != '?'])} 個")
    lines.append(f"- 佐證總數：{total} 件　已完成：{done} 件（{done*100//total if total else 0}%）")

    return "\n".join(lines)


# ════════════════════════════════════════
# 簡述評分表 DOCX 產生
# ════════════════════════════════════════
def build_score_table_docx(db):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    doc = Document()

    # ── 頁面設定 A4 ──
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    def set_font(run, size=11, bold=False, color=None):
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "新細明體"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "新細明體")
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        sz = 14 if level == 1 else 12
        set_font(run, size=sz, bold=True)
        p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def add_cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)

    def shade_cell(cell, hex_color="D9D9D9"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # ── 標題 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("嶺東科技大學教師資格審查簡述評分表（草稿）")
    set_font(r, size=16, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = note.add_run("※ 本表為系統自動產出草稿，敘述文字欄（灰底）請自行補充或修改。")
    set_font(r2, size=9, color=(180, 0, 0))

    doc.add_paragraph()

    # ── 依 section 分組 ──
    current_section = None
    current_title   = None

    for ind_code, ind_info in INDICATORS.items():
        ind_db = db.get(ind_code, {})
        items  = [it for it in ind_db.get("items", []) if it]

        sec   = ind_info["section"]
        title_str = ind_info["title"]
        sub   = ind_info["sub"]
        note_str  = ind_info.get("note", "")

        # 大節標題（壹、教學方面 / 貳、服務方面）
        if sec != current_section:
            current_section = sec
            current_title   = None
            add_heading(sec, level=1)

        # 中節標題（一、教學內容(10%) …）
        if title_str != current_title:
            current_title = title_str
            add_heading(title_str, level=2)

        # ── 子項表格 ──
        # 標頭列
        tbl = doc.add_table(rows=0, cols=4)
        tbl.style = "Table Grid"
        tbl.autofit = False

        # 欄寬：項次 1.5cm、項目 8cm、自評 2cm、行政複評 4.5cm
        widths = [Cm(1.5), Cm(8), Cm(2), Cm(4.5)]
        for i, w in enumerate(widths):
            for cell in tbl.columns[i].cells:
                cell.width = w

        # 標頭
        hdr = tbl.add_row()
        for i, txt in enumerate(["項次", "項目", "自評", "行政單位複評"]):
            shade_cell(hdr.cells[i], "BFBFBF")
            add_cell_text(hdr.cells[i], txt, size=10, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # 子項列
        row = tbl.add_row()
        add_cell_text(row.cells[0], ind_code, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[1], sub, size=10)
        add_cell_text(row.cells[2], str(ind_db.get("self_score", "") or ""),
                      size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(row.cells[3], "", size=10)

        # ── 佐證清單 ──
        ev_row = tbl.add_row()
        # 合併欄
        ev_row.cells[0].merge(ev_row.cells[3])
        shade_cell(ev_row.cells[0], "F2F2F2")

        cell = ev_row.cells[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        label_p = cell.paragraphs[0]
        label_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = label_p.add_run(f"各項次內容表述與佐證：（{ind_code}）")
        set_font(lr, size=9, bold=True, color=(80, 80, 80))

        if note_str:
            np_ = cell.add_paragraph()
            nr = np_.add_run(f"說明：{note_str}")
            set_font(nr, size=8, color=(120, 80, 0))

        if items:
            # 依 group 分組
            groups = {}
            for it in items:
                g = it.get("group", "") or ""
                groups.setdefault(g, []).append(it)

            for g_name, g_items in groups.items():
                if g_name:
                    gp = cell.add_paragraph()
                    gr = gp.add_run(f"▌ {g_name}")
                    set_font(gr, size=9, bold=True)

                for it in g_items:
                    ip = cell.add_paragraph()
                    ip.paragraph_format.left_indent = Cm(0.5)
                    desc = it.get("desc", "")
                    year = it.get("year", "")
                    code = it.get("code", "")
                    line = f"・{desc}"
                    if year:
                        line = f"・[{year}] {desc}"
                    ir = ip.add_run(line)
                    set_font(ir, size=9)
        else:
            ep = cell.add_paragraph()
            er = ep.add_run("（尚無佐證資料）")
            set_font(er, size=9, color=(150, 150, 150))

        # ── 敘述文字欄（留空，待填）──
        narr_row = tbl.add_row()
        narr_row.cells[0].merge(narr_row.cells[3])
        shade_cell(narr_row.cells[0], "FFF2CC")
        nc = narr_row.cells[0]
        np2 = nc.paragraphs[0]
        nr2 = np2.add_run("【敘述文字】請在此補充說明（可貼上已整理好的段落）：")
        set_font(nr2, size=9, bold=True, color=(120, 80, 0))
        # 留幾行空白
        for _ in range(4):
            nc.add_paragraph()

        doc.add_paragraph()  # 指標間距

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ════════════════════════════════════════
# HTTP 伺服器
# ════════════════════════════════════════
class Handler(BaseHTTPRequestHandler):

    def log_message(self, format, *args):
        pass  # 靜音log

    def do_GET(self):
        path = urlparse(self.path).path

        # 首頁
        if path in ("/", "/index.html"):
            content = (STATIC / "index.html").read_bytes()
            self._respond(200, "text/html; charset=utf-8", content)

        # 靜態檔案
        elif path.startswith("/static/"):
            fp = STATIC / path[8:]
            if fp.exists():
                self._respond(200, "application/octet-stream", fp.read_bytes())
            else:
                self._respond(404, "text/plain", b"Not found")

        # 取得資料庫
        elif path == "/api/db":
            db = load_db()
            self._json(db)

        # 取得指標定義
        elif path == "/api/indicators":
            self._json(INDICATORS)

        # 回傳目前設定（供前端顯示 AI 端點）
        elif path == "/api/config":
            self._json({"llama_url": LLAMA_URL})

        # 預覽 upload 檔案（圖片/PDF第一頁）
        elif path.startswith("/api/preview/"):
            from urllib.parse import unquote
            rel = unquote(path[13:])
            fp = UPLOAD / rel
            if not fp.exists():
                # 試著只用純檔名在 uploads 底下搜尋
                fp = next(UPLOAD.rglob(Path(rel).name), None) or fp
            ext = fp.suffix.lower() if fp and fp.exists() else ""
            mime = {
                ".pdf":  "application/pdf",
                ".png":  "image/png",
                ".jpg":  "image/jpeg",
                ".jpeg": "image/jpeg",
                ".bmp":  "image/bmp",
                ".gif":  "image/gif",
                ".webp": "image/webp",
            }.get(ext, "application/octet-stream")
            if fp and fp.exists():
                self.send_response(200)
                self.send_header("Content-Type", mime)
                self.send_header("Content-Length", str(fp.stat().st_size))
                self.send_header("Content-Disposition", "inline")
                self.end_headers()
                self.wfile.write(fp.read_bytes())
            else:
                self._respond(404, "text/plain", b"not found")

        # 下載 output PDF
        elif path.startswith("/api/download/"):
            from urllib.parse import unquote
            fname = unquote(path[14:])
            # 只取純檔名（前端有時會帶路徑）
            fname = Path(fname).name
            fp = OUTPUT / fname
            if fp.exists():
                self.send_response(200)
                self.send_header("Content-Type", "application/pdf")
                # filename* 支援中文
                import urllib.parse
                encoded = urllib.parse.quote(fname, safe="")
                self.send_header("Content-Disposition",
                    f"attachment; filename*=UTF-8''{encoded}")
                self.send_header("Content-Length", str(fp.stat().st_size))
                self.end_headers()
                self.wfile.write(fp.read_bytes())
            else:
                # 列出 output 目錄幫助 debug
                files = [f.name for f in OUTPUT.iterdir()] if OUTPUT.exists() else []
                print(f"[download 404] fname={fname!r}, output有: {files}")
                self._respond(404, "text/plain", b"not found")

        else:
            self._respond(404, "text/plain", b"Not found")

    def do_POST(self):
        path = urlparse(self.path).path
        length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(length)

        # 上傳檔案
        if path == "/api/upload":
            # body = multipart，簡化處理：前端送 JSON {filename, ind_code, data_b64}
            data = json.loads(body)
            filename = data["filename"]
            ind_code = data["ind_code"]
            file_data = base64.b64decode(data["data"])

            # 存到 uploads/{ind_code}/
            folder = UPLOAD / ind_code
            folder.mkdir(parents=True, exist_ok=True)
            save_path = folder / filename
            save_path.write_bytes(file_data)

            # 回傳相對路徑
            rel = f"{ind_code}/{filename}"
            self._json({"ok": True, "path": rel, "pages": self._count_pages(save_path)})

        # 儲存資料庫（整包）
        elif path == "/api/save":
            db = json.loads(body)
            save_db(db)
            self._json({"ok": True})

        # 更新單一指標的 items
        elif path == "/api/update":
            data = json.loads(body)
            db = load_db()
            ind_code = data["ind_code"]
            db[ind_code]["items"] = data["items"]
            if "self_score" in data:
                db[ind_code]["self_score"] = data["self_score"]
            save_db(db)
            self._json({"ok": True})

        # 產出卷夾 PDF
        elif path == "/api/build":
            data = json.loads(body)
            ind_code = data["ind_code"]
            mode     = data.get("mode", "volume")  # volume / divider / stamp
            db = load_db()
            try:
                if mode == "divider":
                    # 只產分隔頁
                    div_bytes = make_divider_bytes(ind_code, db)
                    out_path = OUTPUT / f"{ind_code}_分隔頁.pdf"
                    out_path.write_bytes(div_bytes)
                    self._json({"ok": True, "file": out_path.name})
                elif mode == "stamp":
                    # 只蓋章（不加分隔頁）
                    out_path = build_volume(ind_code, db, include_divider=False)
                    self._json({"ok": True, "file": out_path.name})
                else:
                    # 完整卷夾（預設）
                    out_path = build_volume(ind_code, db, include_divider=True)
                    self._json({"ok": True, "file": out_path.name})
            except Exception as e:
                import traceback
                print(traceback.format_exc())
                self._json({"ok": False, "error": str(e)})

        # 產出全部卷夾（每個指標各一個 PDF，原本功能）
        elif path == "/api/build_all":
            db = load_db()
            built = []
            errors = []
            for ind_code in INDICATORS:
                ind = db.get(ind_code, {})
                if [it for it in ind.get("items",[]) if it]:
                    try:
                        out = build_volume(ind_code, db, include_divider=True)
                        built.append(out.name)
                    except Exception as e:
                        errors.append(f"{ind_code}: {e}")
            self._json({"ok": True, "built": built, "errors": errors})

        # 產出合併卷夾（全部指標合併成一個大 PDF）
        elif path == "/api/build_merge":
            db = load_db()
            merged = fitz.open()
            built = []
            errors = []
            for ind_code in INDICATORS:
                ind = db.get(ind_code, {})
                items = [it for it in ind.get("items", []) if it is not None]
                if not items:
                    continue
                try:
                    tmp_path = build_volume(ind_code, db, include_divider=True)
                    tmp_doc = fitz.open(str(tmp_path))
                    merged.insert_pdf(tmp_doc)
                    tmp_doc.close()
                    built.append(ind_code)
                    print(f"[build_merge] {ind_code} ✅ ({len(items)} 份)")
                except Exception as e:
                    import traceback
                    print(f"[build_merge] {ind_code} ❌ {e}")
                    errors.append(f"{ind_code}: {e}")

            if len(merged) > 0:
                out_path = OUTPUT / "佐證資料完整合併卷夾.pdf"
                merged.save(str(out_path))
                merged.close()
                print(f"[build_merge] 完成！共 {len(built)} 個指標")
                self._json({"ok": True, "file": out_path.name,
                            "built": built, "errors": errors,
                            "total_inds": len(built)})
            else:
                self._json({"ok": False, "error": "沒有任何佐證資料可產出"})

        # AI 批次分析單一指標所有檔案（後端輪詢）
        elif path == "/api/ai_analyze_indicator":
            data = json.loads(body)
            ind_code = data["ind_code"]
            db = load_db()
            ind = db.get(ind_code, {})
            items = ind.get("items", [])
            results = []
            for i, item in enumerate(items):
                file_rel = item.get("file", "")
                if not file_rel:
                    results.append({"idx": i, "ok": False, "error": "無檔案"})
                    continue
                file_path = UPLOAD / file_rel
                if not file_path.exists():
                    results.append({"idx": i, "ok": False, "error": "檔案不存在"})
                    continue
                res = self._llama_analyze(file_path, ind_code)
                if res.get("ok") and res.get("ai"):
                    ai = res["ai"]
                    if ai.get("year"):    items[i]["year"] = ai["year"]
                    if ai.get("summary"): items[i]["desc"] = ai["summary"]
                    if ai.get("ocr_text"):items[i]["ocr"]  = ai["ocr_text"]
                    items[i]["ai"] = ai
                results.append({"idx": i, **res})
            db[ind_code]["items"] = items
            save_db(db)
            self._json({"ok": True, "results": results})

        # AI 分析（送圖給 llama.cpp）
        elif path == "/api/ai_analyze":
            data = json.loads(body)
            file_path = UPLOAD / data["file"]
            result = self._llama_analyze(file_path, data.get("ind_hint",""))
            self._json(result)

        # 重新排序
        elif path == "/api/reorder":
            data = json.loads(body)
            db = load_db()
            ind_code = data["ind_code"]
            new_order = data["order"]  # list of indices
            items = db[ind_code].get("items", [])
            db[ind_code]["items"] = [items[i] for i in new_order if i < len(items)]
            save_db(db)
            self._json({"ok": True})

        # 刪除檔案
        elif path == "/api/delete_file":
            data = json.loads(body)
            fp = UPLOAD / data["file"]
            if fp.exists():
                fp.unlink()
            self._json({"ok": True})

        # 拆頁匯入：把多頁 PDF 拆成每頁，自動加入指標
        elif path == "/api/split_pdf":
            data = json.loads(body)
            ind_code   = data["ind_code"]
            file_rel   = data["file"]        # 已上傳的大 PDF 路徑
            skip_first = data.get("skip_first", True)  # 是否跳過第一頁（目錄）
            start_code = data.get("start_code", "")    # 可選：手動指定流水號起始

            fp = UPLOAD / file_rel
            if not fp.exists():
                self._json({"ok":False,"error":"找不到檔案"}); return

            db = load_db()
            if ind_code not in db:
                db[ind_code] = {**INDICATORS.get(ind_code,{}), "items":[]}

            src = fitz.open(str(fp))
            total = src.page_count
            start_page = 1 if skip_first else 0

            ind_dir = UPLOAD / ind_code
            ind_dir.mkdir(parents=True, exist_ok=True)

            added = []
            import time
            for i in range(start_page, total):
                # 每頁存成獨立 PDF
                page_doc = fitz.open()
                page_doc.insert_pdf(src, from_page=i, to_page=i)
                ts = int(time.time()*1000) + i
                fname = f"split_{ts}_{i+1}.pdf"
                out_path = ind_dir / fname
                page_doc.save(str(out_path))
                page_doc.close()

                rel_path = f"{ind_code}/{fname}"
                item = {
                    "code": "",        # 由 regenCodes 補
                    "year": "",
                    "desc": f"第{i+1}頁",
                    "file": rel_path,
                    "pages": 1,
                    "status": "pending",
                    "ai": None,
                    "ocr": "",
                    "group": ""
                }
                db[ind_code]["items"].append(item)
                added.append(rel_path)

            save_db(db)
            print(f"[split] {ind_code} 拆出 {len(added)} 頁")
            self._json({"ok":True,"added":len(added),"total":total})
        elif path == "/api/move_file":
            data = json.loads(body)
            src = UPLOAD / data["file"]
            to_ind = data["to_ind"]
            dst_dir = UPLOAD / to_ind
            dst_dir.mkdir(parents=True, exist_ok=True)
            dst = dst_dir / src.name
            # 避免同名衝突
            if dst.exists():
                import time
                stem, suf = src.stem, src.suffix
                dst = dst_dir / f"{stem}_{int(time.time())}{suf}"
            if src.exists():
                shutil.move(str(src), str(dst))
            new_rel = f"{to_ind}/{dst.name}"
            self._json({"ok": True, "new_path": new_rel})

        # 智能分類：llama OCR + 後端關鍵字分類（不送指標清單，避免 context 爆）
        elif path == "/api/suggest_indicator":
            import re as _re

            data = json.loads(body)
            file_rel = data.get("file","")
            file_path = UPLOAD / file_rel

            # Step1：llama 只做 OCR（用現有 _llama_analyze）
            ocr_result = self._llama_analyze(file_path, "")
            if not ocr_result.get("ok"):
                self._json({"ok":False,"error":ocr_result.get("error","OCR失敗")}); return

            ai = ocr_result.get("ai",{})
            ocr_text = (ai.get("ocr_text","") or ai.get("summary","")).lower()
            doc_type = ai.get("doc_type","")
            year     = ai.get("year","")
            summary  = ai.get("summary","")
            print(f"[suggest] OCR完成 doc_type={doc_type} year={year}")

            # Step2：後端關鍵字評分（完全不跑第二次 llama）
            IND_KEYWORDS = {
                "A4-5": ["研習","進修","培訓","研討會","工作坊","hours","小時","證明書","研習證明"],
                "B2-4": ["招生","入學","推廣","高中","宣傳","說明會"],
                "B2-1": ["交辦","協助","學校","全校","配合"],
                "B1-1": ["行政","職務","兼任","指派","組長","主任","委員"],
                "B3-2": ["系所","招生","甄選","面試"],
                "B3-1": ["院系","系務","推展","協助"],
                "B4-3": ["導師","社團","指導","球隊","班級"],
                "B4-4": ["輔導","學生","學習"],
                "B5-3": ["社區","校外","服務","義工"],
                "B6-1": ["聘書","聘函","聘期","服務年資","受聘","講師","教授"],
                "B2-5": ["就業","職涯","輔導"],
                "C1-1": ["研究計畫","補助","科技部","國科會","研究案"],
                "C2-1": ["論文","期刊","paper","journal","發表","研討"],
                "C3-1": ["演講","受邀","專題","keynote"],
                "C3-5": ["競賽","獲獎","獎","得獎","金獎","銀獎","佳作"],
                "C3-8": ["展覽","藝文","創作","作品","展出"],
                "C3-10":["顧問","校外","兼職","業界"],
                "D1-1": ["獎狀","獎勵狀","嘉獎","表揚","感謝狀","感謝","績優"],
                "D2-1": ["證照","認證","license","certificate","考試及格"],
                "A6-1": ["教學評量","評量分數","學生評鑑"],
            }

            scores = {}
            for code, kws in IND_KEYWORDS.items():
                if code not in INDICATORS: continue
                s = sum(2 if kw in ocr_text else 0 for kw in kws)
                # doc_type 加權
                if doc_type in ("獎勵狀","獎狀","感謝狀") and code=="D1-1": s+=3
                if doc_type in ("研習證明","進修證明","研習證書") and code=="A4-5": s+=5
                if doc_type == "聘書" and code=="B6-1": s+=4
                if s>0: scores[code]=s

            if scores:
                best = max(scores, key=lambda c:scores[c])
                conf = "high" if scores[best]>=4 else "medium" if scores[best]>=2 else "low"
            else:
                best = "?"
                conf = "low"

            print(f"[suggest] 關鍵字分類：{best} ({conf}) scores={sorted(scores.items(),key=lambda x:-x[1])[:5]}")
            self._json({"ok":True,"suggestion":{
                "ind_code": best,
                "reason": summary[:30] if summary else doc_type,
                "year": year,
                "confidence": conf,
                "doc_type": doc_type,
                "ocr": ai.get("ocr_text","")[:200],
            }})

        # 匯出精華摘要 PDF（分隔頁 + skip_pages 摘要頁）
        elif path == "/api/export_summary_pdf":
            try:
                db = load_db()
                merged = fitz.open()
                collected = []
                errors = []

                for ind_code in INDICATORS:
                    ind = db.get(ind_code, {})
                    items = [it for it in ind.get("items", []) if it]
                    for it in items:
                        skip = it.get("skip_pages") or 0
                        if skip <= 0:
                            continue
                        file_rel = it.get("file", "")
                        if not file_rel:
                            continue
                        file_path = UPLOAD / file_rel
                        if not file_path.exists():
                            errors.append(f"{it.get('code','')} 檔案不存在: {file_rel}")
                            continue

                        try:
                            # 1. 先加分隔頁（展開模式）
                            div_bytes = make_divider_bytes(ind_code, db, expanded=True)
                            div_doc = fitz.open("pdf", div_bytes)
                            merged.insert_pdf(div_doc)
                            div_doc.close()

                            # 2. 加 skip_pages 頁（摘要/清單頁）
                            ext = file_path.suffix.lower()
                            IMG_EXTS = {".jpg",".jpeg",".png",".bmp",".gif",".webp"}
                            if ext in IMG_EXTS:
                                img_bytes = file_path.read_bytes()
                                content_bytes = img_to_pdf_bytes(img_bytes, ext.lstrip("."))
                            else:
                                content_bytes = file_path.read_bytes()

                            src_doc = fitz.open("pdf", content_bytes)
                            pages_to_take = min(skip, len(src_doc))
                            merged.insert_pdf(src_doc, from_page=0, to_page=pages_to_take - 1)
                            src_doc.close()

                            collected.append(f"{it.get('code','')} skip={skip}")
                            print(f"[summary] ✅ {it.get('code','')} skip={skip} {file_rel}")
                        except Exception as e:
                            errors.append(f"{it.get('code','')}: {e}")
                            print(f"[summary] ❌ {it.get('code','')} {e}")

                if len(merged) > 0:
                    out_path = OUTPUT / "簡述評分表_精華摘要.pdf"
                    merged.save(str(out_path))
                    merged.close()
                    out_bytes = out_path.read_bytes()
                    print(f"[summary] 完成！共 {len(collected)} 個項目，{len(merged) if False else '?'} 頁")
                    self._respond(200, "application/octet-stream", out_bytes)
                else:
                    merged.close()
                    self._json({"ok": False, "error": "沒有找到任何 skip_pages 項目", "errors": errors})

            except Exception as e:
                import traceback
                print(traceback.format_exc())
                self._json({"ok": False, "error": str(e)})

        # 匯出簡述評分表 docx
        elif path == "/api/export_score_table":
            try:
                db = load_db()
                docx_bytes = build_score_table_docx(db)
                out_path = OUTPUT / "教師資格審查簡述評分表_草稿.docx"
                out_path.write_bytes(docx_bytes)
                self._respond(200, "application/octet-stream", docx_bytes)
            except Exception as e:
                import traceback
                print(traceback.format_exc())
                self._json({"ok": False, "error": str(e)})

        # ── 匯出 GPTs Knowledge 檔 ──────────────────────────────
        elif path == "/api/export_gpts_knowledge":
            try:
                db = load_db()
                md_text = build_gpts_knowledge(db)
                md_bytes = md_text.encode("utf-8")
                self.send_response(200)
                self.send_header("Content-Type", "application/octet-stream")
                self.send_header("Content-Length", str(len(md_bytes)))
                self.send_header("Content-Disposition",
                                 "attachment; filename*=UTF-8''gpts_knowledge_full.md")
                self.send_header("Access-Control-Allow-Origin", "*")
                self.end_headers()
                self.wfile.write(md_bytes)
            except Exception as e:
                import traceback
                print(traceback.format_exc())
                self._json({"ok": False, "error": str(e)})

        else:
            self._respond(404, "text/plain", b"Not found")

    def _llama_analyze(self, file_path, ind_hint=""):
        """呼叫 llama.cpp Vision 分析文件"""
        try:
            # PDF 第一頁轉圖片，圖片直接讀
            ext = file_path.suffix.lower()
            IMG_EXTS = {".jpg",".jpeg",".png",".bmp",".gif",".webp"}
            if ext == ".pdf":
                doc = fitz.open(str(file_path))
                pix = doc[0].get_pixmap(dpi=150)
                img_bytes = pix.tobytes("jpeg")
            elif ext in IMG_EXTS:
                img_bytes = file_path.read_bytes()
            else:
                return {"ok": False, "error": f"不支援的格式：{ext}"}

            b64 = base64.b64encode(img_bytes).decode("ascii")

            # ── Prompt：簡潔、直接要 JSON ──
            # /no_think 要放在 system 訊息，不是 user
            hint_line = f"（此文件對應評鑑指標：{ind_hint}）" if ind_hint else ""
            prompt = f"""請看這張台灣大學教師升等佐證文件圖片。{hint_line}

請直接輸出以下 JSON，不要任何說明或 markdown：
{{"doc_type":"文件類型(獎狀/聘書/感謝狀/邀請函/派令/考績表/其他)","org":"發文機構","year":"學年度數字或範圍如104-114","summary":"25字內說明","ocr_text":"圖片全部文字","confidence":"high/medium/low"}}"""

            req_body = {
                "model": "qwen",
                "max_tokens": 2048,
                "temperature": 0.0,
                "messages": [
                    {
                        "role": "system",
                        "content": "你是文件分析助理。只輸出 JSON，不思考，不解釋。/no_think"
                    },
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
                            },
                            {
                                "type": "text",
                                "text": prompt
                            }
                        ]
                    }
                ]
            }

            req_data = json.dumps(req_body, ensure_ascii=False).encode("utf-8")
            req = urllib.request.Request(
                f"{LLAMA_URL}/v1/chat/completions",
                data=req_data,
                headers={"Content-Type": "application/json; charset=utf-8"},
            )
            with urllib.request.urlopen(req, timeout=90) as resp:
                result = json.loads(resp.read().decode("utf-8"))

            # 取回應文字
            raw = result["choices"][0]["message"]["content"]
            print(f"[AI raw] {raw[:400]}")

            # 移除 <think>...</think> 區塊（Qwen3 有時還是會輸出）
            import re
            raw = re.sub(r"<think>[\s\S]*?</think>", "", raw).strip()

            # 找 JSON：先找 ```json 區塊，再找裸 {}
            m = re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
            json_str = m.group(1).strip() if m else raw

            # 找最外層 {}
            m2 = re.search(r"\{[\s\S]*\}", json_str)
            if m2:
                try:
                    parsed = json.loads(m2.group())
                    return {"ok": True, "ai": parsed}
                except json.JSONDecodeError as je:
                    print(f"[AI JSON parse error] {je}")
                    print(f"[AI raw json_str] {json_str[:300]}")

            # fallback：把原始文字當 OCR 全文
            return {"ok": True, "ai": {
                "doc_type": "未識別",
                "org": "",
                "year": None,
                "summary": raw[:40],
                "ocr_text": raw,
                "confidence": "low"
            }}

        except Exception as e:
            return {"ok": False, "error": str(e)}

    def _count_pages(self, path):
        try:
            pass  # fitz already imported at top
            if path.suffix.lower() == ".pdf":
                return len(fitz.open(str(path)))
        except:
            pass
        return 1

    def _respond(self, code, ctype, body):
        self.send_response(code)
        self.send_header("Content-Type", ctype)
        self.send_header("Content-Length", str(len(body)))
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(body)

    def _json(self, data):
        body = json.dumps(data, ensure_ascii=False).encode()
        self._respond(200, "application/json; charset=utf-8", body)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()


if __name__ == "__main__":
    PORT = 8899
    server = HTTPServer(("localhost", PORT), Handler)
    print(f"""
+--------------------------------------+
|  Server started                      |
|  http://localhost:{PORT}              |
|  Press Ctrl+C to stop                |
+--------------------------------------+
""")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n系統已停止")